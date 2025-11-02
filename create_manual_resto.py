#!/usr/bin/env python3
"""
Script to create comprehensive user manual for Ladapala Restaurant POS System
in Bahasa Indonesia
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Set document title
    title = doc.add_heading('PANDUAN PENGGUNAAN SISTEM POS RESTORAN LADAPALA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version info
    version = doc.add_paragraph('Versi 1.0 - November 2025')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('DAFTAR ISI', 1)
    toc_items = [
        '1. Pendahuluan',
        '2. Memulai Sistem',
        '   2.1. Login ke Sistem',
        '   2.2. Buka Shift Kasir',
        '   2.3. Mengenal Dashboard',
        '3. Operasional Kasir',
        '   3.1. Dashboard Kasir',
        '   3.2. Membuat Pesanan (Menu)',
        '   3.3. Manajemen Meja',
        '   3.4. Proses Pembayaran (Transaksi)',
        '   3.5. Tutup Shift Kasir',
        '4. Office - Manajemen & Administrasi',
        '   4.1. Office Dashboard',
        '   4.2. Manajemen Stok',
        '   4.3. Manajemen Vendor',
        '   4.4. Resep & Menu',
        '   4.5. Riwayat Penjualan',
        '   4.6. Data Pelanggan',
        '   4.7. Jadwal Karyawan',
        '   4.8. Laporan',
        '   4.9. Pengaturan',
        '5. Dapur (Kitchen)',
        '   5.1. Dashboard Dapur',
        '   5.2. Menerima Pesanan',
        '   5.3. Update Status Pesanan',
        '6. Fitur Umum',
        '   6.1. Profil Pengguna',
        '   6.2. Pengaturan',
        '   6.3. Logout',
        '7. Tips & Troubleshooting',
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('1. PENDAHULUAN', 1)

    doc.add_heading('Tentang Sistem', 2)
    doc.add_paragraph(
        'Sistem POS (Point of Sale) Restoran Ladapala adalah platform terintegrasi yang dirancang '
        'khusus untuk mengelola operasional restoran Indonesia. Sistem ini mencakup manajemen pesanan, '
        'pembayaran, stok, shift kasir, laporan penjualan, dan koordinasi dapur dalam satu platform '
        'yang mudah digunakan.'
    )

    doc.add_heading('Dua Area Utama Sistem', 2)

    areas = [
        ('Kasir (Dashboard Hijau)',
         'Area untuk operasional harian kasir seperti menerima pesanan, manajemen meja, '
         'proses pembayaran, dan tutup shift. Warna hijau neon (#58ff34) menandakan area ini.'),
        ('Office (Dashboard Biru)',
         'Area untuk manajemen dan administrasi seperti stok, vendor, resep, laporan penjualan, '
         'jadwal karyawan, dan pengaturan sistem. Warna biru menandakan area ini.'),
    ]

    for area_name, description in areas:
        p = doc.add_paragraph()
        p.add_run(area_name).bold = True
        p.add_run(f'\n{description}')

    doc.add_heading('Hak Akses Berdasarkan Role', 2)
    doc.add_paragraph(
        'Setiap pengguna memiliki hak akses berbeda berdasarkan role mereka:'
    )

    access_list = [
        'ADMIN: Akses penuh ke semua area (Kasir dan Office)',
        'MANAGER: Akses ke Office dan dapat override jadwal kasir',
        'CASHIER: Akses ke area Kasir (buka/tutup shift, pesanan, pembayaran)',
        'KITCHEN: Akses ke dashboard dapur (melihat dan update pesanan)',
        'WAREHOUSE: Akses ke manajemen stok',
    ]

    for access in access_list:
        doc.add_paragraph(access, style='List Bullet')

    doc.add_heading('Konsep Penting: Order Status Flow', 2)
    doc.add_paragraph('Pesanan di sistem mengikuti alur status yang mencerminkan operasional restoran:')

    status_flow = [
        'CONFIRMED: Pesanan diterima dan dikonfirmasi',
        'PREPARING: Dapur sedang memasak pesanan',
        'READY: Makanan sudah siap, menunggu diantar ke pelanggan',
        'COMPLETED: Makanan sudah diantar ke meja pelanggan, siap untuk pembayaran',
        'CANCELLED: Pesanan dibatalkan',
    ]

    for status in status_flow:
        doc.add_paragraph(status, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('PENTING: ').bold = True
    p.add_run(
        'Hanya pesanan dengan status COMPLETED yang bisa diproses untuk pembayaran. '
        'READY berarti makanan siap di dapur (belum diantar), sedangkan COMPLETED berarti '
        'makanan sudah diantar ke meja pelanggan.'
    )

    doc.add_page_break()

    # Chapter 2: Getting Started
    doc.add_heading('2. MEMULAI SISTEM', 1)

    doc.add_heading('2.1. Login ke Sistem', 2)

    doc.add_paragraph('Langkah-langkah untuk login:')

    login_steps = [
        'Buka browser (Chrome, Firefox, Safari, atau Edge)',
        'Ketik alamat: http://localhost:3000/login',
        'Anda akan melihat halaman login dengan background hijau neon',
        'Masukkan email Anda di kolom "Email"',
        'Masukkan password Anda di kolom "Password"',
        'Klik tombol "Masuk" berwarna hijau neon',
        'Jika berhasil, Anda akan diarahkan ke dashboard sesuai role Anda',
    ]

    for i, step in enumerate(login_steps, 1):
        doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Catatan Penting:', style='Heading 3')
    notes = [
        'Pastikan Anda menggunakan email dan password yang telah diberikan oleh administrator',
        'Jika lupa password, hubungi administrator',
        'Sistem akan mengingat sesi Anda, jadi tidak perlu login ulang setiap kali',
        'Untuk keamanan, selalu logout setelah selesai bekerja',
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_heading('2.2. Buka Shift Kasir', 2)

    doc.add_paragraph(
        'Sebelum mulai menerima pesanan dan pembayaran, kasir HARUS membuka shift terlebih dahulu. '
        'Ini adalah langkah wajib untuk tracking penjualan dan cash reconciliation.'
    )

    doc.add_paragraph('Langkah-langkah Buka Shift:', style='Heading 3')

    open_shift = [
        'Setelah login, jika belum ada shift aktif, Anda akan melihat notifikasi "Belum ada shift aktif"',
        'Klik menu hamburger (☰) di sidebar bawah',
        'Pilih "Buka Shift"',
        'Anda akan diarahkan ke halaman Buka Shift Kasir',
        'Isi form buka shift:',
        '   - Shift: Pilih shift Anda (Pagi/Siang/Sore/Malam)',
        '   - Modal Awal: Masukkan jumlah uang tunai awal di laci kasir',
        '   - Catatan: Opsional, untuk catatan khusus',
        'Sistem akan validasi jadwal Anda:',
        '   - Jika sesuai jadwal: Shift langsung terbuka',
        '   - Jika tidak sesuai jadwal: Perlu approval Manager (override)',
        'Klik "Buka Shift"',
        'Setelah berhasil, Anda akan diarahkan ke Dashboard Kasir',
        'Status shift aktif akan muncul di top navbar',
    ]

    for i, step in enumerate(open_shift, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Tips Buka Shift:', style='Heading 3')
    shift_tips = [
        'Hitung modal awal dengan teliti, akan dibandingkan saat tutup shift',
        'Jika shift Anda tidak sesuai jadwal, hubungi Manager untuk approval',
        'Hanya satu shift aktif per kasir dalam satu waktu',
        'Semua transaksi akan tercatat di shift yang aktif',
    ]
    for tip in shift_tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('2.3. Mengenal Dashboard', 2)

    doc.add_paragraph(
        'Setelah shift dibuka, Anda akan melihat Dashboard Kasir dengan warna hijau neon. '
        'Dashboard yang muncul tergantung pada role Anda.'
    )

    doc.add_paragraph('Elemen-elemen Dashboard Kasir:', style='Heading 3')

    doc.add_paragraph('1. Sidebar (Kolom Kiri):')
    sidebar_items = [
        'Ikon toko/restoran di bagian atas',
        'Menu utama: Beranda, Menu, Meja, Transaksi',
        'Tombol menu hamburger (☰) di bawah untuk akses:',
        '   - Office (jika punya akses)',
        '   - Profile',
        '   - Settings',
        '   - Buka Shift / Tutup Shift',
        '   - Logout',
        'Warna hijau neon menandakan menu aktif',
    ]
    for item in sidebar_items:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('2. Top Navbar (Bagian Atas):')
    navbar_items = [
        'Status Shift: Menampilkan shift aktif dan durasi (contoh: "SHIFT PAGI - 02:45:30")',
        'Breadcrumb: Menunjukkan lokasi Anda di sistem',
        'Nama pengguna dan role di pojok kanan',
        'Badge merah untuk notifikasi (jika ada)',
    ]
    for item in navbar_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('3. Konten Utama (Tengah):')
    doc.add_paragraph(
        'Area ini menampilkan statistik operasional, daftar pesanan, '
        'atau form input sesuai dengan menu yang dipilih.'
    )

    doc.add_page_break()

    # Chapter 3: Cashier Operations
    doc.add_heading('3. OPERASIONAL KASIR', 1)

    doc.add_heading('3.1. Dashboard Kasir', 2)

    doc.add_paragraph(
        'Dashboard Kasir adalah halaman utama untuk kasir. Di sini Anda dapat melihat '
        'ringkasan operasional shift Anda.'
    )

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Beranda" (rumah) di sidebar')

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    dashboard_info = [
        'Kartu Statistik Shift:',
        '   - Total Penjualan Shift Ini: Akumulasi penjualan sejak shift dibuka',
        '   - Total Pesanan: Jumlah pesanan yang diproses',
        '   - Meja Terisi: Jumlah meja yang sedang occupied',
        '   - Pesanan Menunggu: Pesanan yang perlu diproses',
        'Pesanan Terbaru: Daftar pesanan terbaru dengan status',
        'Meja Aktif: Daftar meja yang sedang terisi dengan info ringkas',
        'Quick Actions: Tombol cepat untuk aksi umum',
    ]
    for item in dashboard_info:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2. Membuat Pesanan (Menu)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Menu" (menu restoran) di sidebar')

    doc.add_paragraph('Fungsi halaman ini:', style='Heading 3')
    doc.add_paragraph(
        'Membuat pesanan baru dengan memilih menu dari katalog produk restoran.'
    )

    doc.add_paragraph('A. Tampilan Halaman Menu', style='Heading 3')
    menu_layout = [
        'Layout 2 kolom:',
        '   Kiri: Katalog menu dengan kategori',
        '   Kanan: Keranjang pesanan (order cart)',
        'Filter kategori: All, Main Course, Appetizer, Beverage, Dessert, dll',
        'Search bar: Cari menu berdasarkan nama',
        'Kartu menu menampilkan: Foto, Nama, Harga, Stok tersedia',
    ]
    for item in menu_layout:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Membuat Pesanan Baru', style='Heading 3')
    doc.add_paragraph('Langkah-langkah:')

    create_order = [
        'Pilih Meja atau Pilih "Takeaway" (jika bawa pulang)',
        'Jika Meja: Pilih nomor meja dari dropdown',
        'Browse menu atau gunakan search untuk mencari',
        'Klik kartu menu untuk menambahkan ke keranjang',
        'Di keranjang (kolom kanan):',
        '   - Lihat item yang sudah dipilih',
        '   - Adjust quantity dengan tombol + / -',
        '   - Hapus item dengan tombol X',
        '   - Lihat subtotal per item dan total keseluruhan',
        'Tambahkan catatan khusus jika ada (opsional)',
        'Contoh catatan: "Tidak pedas", "Extra sambal", "Tanpa bawang"',
        'Klik "Buat Pesanan" untuk konfirmasi',
        'Pesanan tersimpan dengan status CONFIRMED',
        'Dapur akan menerima notifikasi pesanan baru',
    ]

    for i, step in enumerate(create_order, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Tips Membuat Pesanan', style='Heading 3')
    order_tips = [
        'Pastikan meja dipilih dengan benar sebelum buat pesanan',
        'Cek stok tersedia sebelum menambahkan item (jika stok 0, tidak bisa dipesan)',
        'Double check pesanan sebelum klik "Buat Pesanan"',
        'Untuk pesanan takeaway, pilih "Takeaway" bukan meja',
        'Gunakan catatan untuk permintaan khusus pelanggan',
    ]
    for tip in order_tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('3.3. Manajemen Meja', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Meja" (cells/grid) di sidebar')

    doc.add_paragraph('Fungsi halaman ini:', style='Heading 3')
    doc.add_paragraph(
        'Memonitor status semua meja restoran dan pesanan di setiap meja. '
        'Halaman ini dirancang untuk display publik, jadi tidak menampilkan harga/revenue.'
    )

    doc.add_paragraph('Status Meja:', style='Heading 3')
    table_statuses = [
        'AVAILABLE (Hijau): Meja kosong, siap digunakan',
        'OCCUPIED (Biru): Meja sedang digunakan pelanggan',
        'RESERVED (Ungu): Meja sudah dipesan (via localStorage)',
    ]
    for status in table_statuses:
        doc.add_paragraph(status, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Informasi di Kartu Meja (untuk meja OCCUPIED):', style='Heading 3')
    table_info = [
        'Nomor meja',
        'Jumlah pesanan aktif',
        'Status pesanan terakhir (badge berwarna)',
        'Durasi meja terisi',
        'TIDAK menampilkan harga/total (untuk privacy)',
    ]
    for info in table_info:
        doc.add_paragraph(info, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Menggunakan Halaman Meja:', style='Heading 3')
    use_table = [
        'Lihat sekilas meja mana yang tersedia',
        'Monitor status pesanan per meja (badge warna)',
        'Klik kartu meja untuk detail lengkap (popup)',
        'Di popup detail (hanya untuk staff):',
        '   - Lihat semua pesanan di meja tersebut',
        '   - Lihat total tagihan',
        '   - Update status pesanan',
        '   - Proses pembayaran (redirect ke Transaksi)',
    ]
    for i, item in enumerate(use_table, 1):
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{item}', style='List Number')

    doc.add_heading('3.4. Proses Pembayaran (Transaksi)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Transaksi" (credit card) di sidebar')

    doc.add_paragraph('Fungsi halaman ini:', style='Heading 3')
    doc.add_paragraph(
        'Memproses pembayaran untuk pesanan yang sudah COMPLETED (makanan sudah diantar ke pelanggan).'
    )

    doc.add_paragraph('PENTING:', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('Hanya pesanan dengan status COMPLETED yang muncul di halaman transaksi. ').bold = True
    p.add_run(
        'Pastikan dapur sudah mengupdate status pesanan menjadi READY (siap), '
        'lalu waiter/kasir update menjadi COMPLETED (sudah diantar) sebelum proses pembayaran.'
    )

    doc.add_paragraph()
    doc.add_paragraph('A. Melihat Pesanan Siap Bayar', style='Heading 3')
    ready_payment = [
        'Halaman menampilkan daftar pesanan dengan status COMPLETED',
        'Dikelompokkan berdasarkan meja (atau Takeaway)',
        'Informasi yang ditampilkan:',
        '   - Nomor meja',
        '   - Daftar item pesanan (nama + quantity)',
        '   - Total tagihan (sudah termasuk tax jika ada)',
        '   - Durasi pesanan',
        'Filter berdasarkan meja atau search',
    ]
    for item in ready_payment:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Proses Pembayaran', style='Heading 3')
    doc.add_paragraph('Langkah-langkah:')

    process_payment = [
        'Klik pada pesanan yang akan dibayar',
        'Modal pembayaran akan muncul dengan rincian:',
        '   - List semua item dengan harga',
        '   - Subtotal',
        '   - Tax (jika dikonfigurasi)',
        '   - Grand Total',
        'Pilih Metode Pembayaran:',
        '   - CASH (Tunai)',
        '   - CARD (Kartu Debit/Kredit)',
        '   - QRIS',
        '   - TRANSFER',
        'Untuk pembayaran CASH:',
        '   - Masukkan jumlah uang yang diterima dari pelanggan',
        '   - Sistem otomatis hitung kembalian',
        '   - Pastikan kembalian benar sebelum lanjut',
        'Untuk pembayaran NON-CASH:',
        '   - Tidak perlu input jumlah (otomatis sama dengan total)',
        'Klik "Proses Pembayaran"',
        'Konfirmasi jika diminta',
        'Pembayaran berhasil! Modal sukses akan muncul',
        'Struk otomatis ter-generate dan bisa di-print',
        'Status pesanan berubah menjadi PAID',
        'Meja akan berubah status menjadi AVAILABLE (jika semua pesanan di meja sudah bayar)',
        'Transaksi tercatat di shift kasir yang aktif',
    ]

    for i, step in enumerate(process_payment, 1):
        if step.startswith('   -'):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        elif step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Print Struk', style='Heading 3')
    print_receipt = [
        'Setelah pembayaran sukses, modal akan menampilkan preview struk',
        'Struk berisi:',
        '   - Nama restoran dan alamat',
        '   - Nomor pesanan dan tanggal',
        '   - Nama kasir',
        '   - List item dengan harga',
        '   - Subtotal, tax, total',
        '   - Metode pembayaran',
        '   - Uang diterima dan kembalian (untuk cash)',
        '   - Footer dengan ucapan terima kasih',
        'Klik "Print" untuk print struk (akan membuka print dialog)',
        'Atau klik "Close" untuk menutup tanpa print',
    ]
    for item in print_receipt:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('D. Split Bill (Bayar Terpisah)', style='Heading 3')
    doc.add_paragraph(
        'Jika pelanggan ingin bayar terpisah per item, sistem sudah mendukung multiple payments '
        'untuk satu meja. Proses pembayaran dapat dilakukan beberapa kali untuk pesanan yang berbeda '
        'di meja yang sama.'
    )

    doc.add_heading('3.5. Tutup Shift Kasir', 2)

    doc.add_paragraph(
        'Setelah selesai jam kerja, kasir HARUS menutup shift untuk reconciliation (mencocokkan) '
        'uang tunai dan generate laporan shift.'
    )

    doc.add_paragraph('PENTING Sebelum Tutup Shift:', style='Heading 3')
    before_close = [
        'Pastikan SEMUA pesanan di shift ini sudah dibayar (settled)',
        'Sistem akan mencegah tutup shift jika masih ada pesanan unpaid',
        'Check semua meja sudah AVAILABLE',
        'Siapkan uang tunai di laci kasir untuk dihitung',
    ]
    for item in before_close:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Langkah-langkah Tutup Shift:', style='Heading 3')

    close_shift = [
        'Klik menu hamburger (☰) di sidebar',
        'Pilih "Tutup Shift"',
        'Sistem akan validasi semua pesanan sudah settled',
        'Jika ada pesanan unpaid: Sistem akan warning dan tidak bisa tutup shift',
        'Jika semua OK: Halaman tutup shift akan muncul',
        'Lihat Summary Shift:',
        '   - Shift yang aktif (Pagi/Siang/Sore/Malam)',
        '   - Durasi shift',
        '   - Total transaksi',
        '   - Total penjualan',
        'Cash Reconciliation:',
        '   - Expected Cash: Uang yang seharusnya ada (modal awal + pembayaran cash)',
        '   - Actual Cash: Input jumlah uang tunai yang benar-benar ada di laci',
        '   - Cash Difference: Sistem otomatis hitung selisih',
        'Jika ada selisih (difference):',
        '   - Merah: Kurang (shortage)',
        '   - Hijau: Lebih (overage)',
        '   - Selisih > Rp 1,000 akan dicatat di audit log',
        'Breakdown Metode Pembayaran:',
        '   - Cash: Jumlah dan total',
        '   - Card: Jumlah dan total',
        '   - QRIS: Jumlah dan total',
        '   - Transfer: Jumlah dan total',
        'Tambahkan catatan penutupan (opsional)',
        'Contoh: "Selisih karena customer kurang bayar Rp 500"',
        'Klik "Tutup Shift"',
        'Sistem akan:',
        '   - Close shift di database',
        '   - Generate laporan shift (Settlement Report)',
        '   - Log audit trail',
        '   - Reset status shift aktif',
        'Download atau Print Settlement Report',
        'Serahkan laporan dan uang tunai ke supervisor/manager',
    ]

    for i, step in enumerate(close_shift, 1):
        if step.startswith('   -'):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        elif step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Tips Tutup Shift:', style='Heading 3')
    close_tips = [
        'Hitung uang tunai dengan teliti sebelum input Actual Cash',
        'Jika ada selisih besar, cek ulang perhitungan',
        'Simpan Settlement Report untuk arsip',
        'Jangan tutup shift jika masih ada pelanggan di meja',
        'Koordinasi dengan shift berikutnya untuk serah terima',
    ]
    for tip in close_tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_page_break()

    # Chapter 4: Office Operations
    doc.add_heading('4. OFFICE - MANAJEMEN & ADMINISTRASI', 1)

    doc.add_heading('4.1. Office Dashboard', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph(
        'Klik menu hamburger (☰) di sidebar → Pilih "Office" '
        '(hanya tersedia untuk Admin, Manager, dan Warehouse)'
    )

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Dashboard Office adalah pusat kontrol untuk manajemen dan administrasi restoran. '
        'Menampilkan metrics penting dan akses ke semua modul administrasi.'
    )

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    office_dashboard = [
        'Sales Overview: Total penjualan hari ini, minggu ini, bulan ini',
        'Orders Statistics: Total pesanan dan breakdown per status',
        'Stock Alerts: Item yang stok-nya rendah atau hampir expired',
        'Top Selling Items: Menu terlaris',
        'Revenue Chart: Grafik penjualan per periode',
        'Recent Transactions: Transaksi terbaru',
    ]
    for item in office_dashboard:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.2. Manajemen Stok', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Dari Office → Menu "Stok" di sidebar Office')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola inventory bahan baku dan produk restoran termasuk stock in/out, '
        'transfer antar gudang, purchase orders, dan monitoring expiry date.'
    )

    doc.add_paragraph('Modul-modul Stok:', style='Heading 3')

    doc.add_paragraph('A. Stock Items (Master Data)', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Items')
    stock_items = [
        'Daftar semua item di inventory',
        'Kategori: Raw Material, Finished Product, Beverage, Seasoning, Packaging',
        'Informasi per item:',
        '   - Nama item',
        '   - Kategori',
        '   - Unit (kg, liter, pcs, dll)',
        '   - Stock saat ini',
        '   - Minimum stock (alert threshold)',
        '   - Harga per unit',
        'Tambah item baru: Klik "+ Add Item"',
        'Edit/Delete item: Klik ikon aksi di baris item',
        'Alert merah: Item dengan stok di bawah minimum',
    ]
    for item in stock_items:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Stock Movements', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Movements')
    stock_movements = [
        'History semua pergerakan stok (in/out/transfer/adjustment)',
        'Filter berdasarkan:',
        '   - Jenis movement: IN, OUT, TRANSFER, ADJUSTMENT',
        '   - Date range',
        '   - Item',
        'Export data ke Excel untuk analisis',
    ]
    for item in stock_movements:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('C. Stock Receipt (Terima Barang)', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Receipt')
    doc.add_paragraph('Cara mencatat stock in:')
    stock_receipt = [
        'Klik "+ New Receipt"',
        'Pilih item yang diterima',
        'Input quantity',
        'Input harga (opsional)',
        'Pilih vendor/supplier (opsional)',
        'Input expiry date (untuk perishable items)',
        'Tambahkan catatan (nomor invoice, dll)',
        'Klik "Save"',
        'Stock otomatis bertambah di sistem',
    ]
    for step in stock_receipt:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('D. Stock Update (Manual)', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Update')
    doc.add_paragraph(
        'Untuk update stok manual saat produksi atau penggunaan harian. '
        'Pilih item, input quantity yang digunakan, dan sistem akan kurangi stok otomatis.'
    )

    doc.add_paragraph()
    doc.add_paragraph('E. Stock Transfer', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Transfer')
    doc.add_paragraph(
        'Untuk transfer stok antar lokasi (contoh: dari gudang utama ke gudang cabang). '
        'Input item, quantity, lokasi asal, dan lokasi tujuan.'
    )

    doc.add_paragraph()
    doc.add_paragraph('F. Stock Adjustment', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Adjustment')
    doc.add_paragraph(
        'Untuk koreksi stok saat stock opname (hasil perhitungan fisik berbeda dengan sistem). '
        'Input stok aktual, sistem akan adjust otomatis dan catat selisihnya.'
    )

    doc.add_paragraph()
    doc.add_paragraph('G. Expiry Monitoring', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Expiry')
    expiry_monitoring = [
        'Daftar item yang punya expiry date',
        'Badge warna:',
        '   - Merah: Expired atau < 7 hari',
        '   - Kuning: 7-30 hari',
        '   - Hijau: > 30 hari',
        'Sort berdasarkan expiry date terdekat',
        'Action: Tandai sebagai expired/disposal',
    ]
    for item in expiry_monitoring:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('H. Purchase Orders', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Purchase Orders')
    purchase_orders = [
        'Buat PO untuk pemesanan barang ke vendor',
        'Cara membuat PO:',
        '   1. Klik "+ Create PO"',
        '   2. Pilih vendor',
        '   3. Tambahkan items yang akan dibeli',
        '   4. Input quantity dan harga',
        '   5. Expected delivery date',
        '   6. Klik "Create"',
        'Status PO:',
        '   - DRAFT: Masih draft',
        '   - SENT: Sudah dikirim ke vendor',
        '   - PARTIAL: Barang diterima sebagian',
        '   - RECEIVED: Semua barang diterima',
        '   - CANCELLED: Dibatalkan',
        'Receive PO: Update status saat barang datang',
        'Sistem otomatis update stok saat PO received',
    ]
    for i, item in enumerate(purchase_orders):
        if item.startswith('   -'):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        elif item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Number')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('I. Stock Reports', style='Heading 4')
    doc.add_paragraph('Akses: Stok → Reports')
    stock_reports = [
        'Stock Position Report: Current stock semua item',
        'Stock Movement Report: History pergerakan per periode',
        'Low Stock Report: Item di bawah minimum stock',
        'Expiry Report: Item yang akan expired',
        'Stock Value Report: Nilai inventory',
        'Export semua laporan ke Excel/PDF',
    ]
    for item in stock_reports:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.3. Manajemen Vendor', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Vendor" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph('Mengelola database vendor/supplier yang memasok bahan baku ke restoran.')

    doc.add_paragraph('Fitur-fitur:', style='Heading 3')
    vendor_features = [
        'Daftar semua vendor',
        'Informasi vendor:',
        '   - Nama perusahaan',
        '   - Contact person',
        '   - Email, telepon',
        '   - Alamat',
        '   - Kategori produk yang disupply',
        '   - Payment terms',
        'Tambah vendor baru: "+ Add Vendor"',
        'Edit/Delete vendor',
        'Lihat riwayat purchase orders ke vendor',
        'Vendor rating/performance (jika dikonfigurasi)',
    ]
    for item in vendor_features:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.4. Resep & Menu', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Recipe" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola resep menu (recipe) dengan breakdown bahan baku yang dibutuhkan. '
        'Sistem akan auto-deduct stok saat pesanan dibuat berdasarkan recipe.'
    )

    doc.add_paragraph('A. Daftar Resep', style='Heading 3')
    recipe_list = [
        'Semua menu yang punya resep',
        'Informasi: Nama menu, Kategori, Portion size, Cost per portion',
        'Klik menu untuk lihat detail resep',
    ]
    for item in recipe_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Membuat/Edit Resep', style='Heading 3')
    create_recipe = [
        'Pilih menu',
        'Klik "Add Ingredients"',
        'Tambahkan bahan baku yang dibutuhkan:',
        '   - Pilih stock item',
        '   - Input quantity',
        '   - Unit',
        'Sistem akan hitung cost per portion otomatis',
        'Save recipe',
        'Saat pesanan dibuat, sistem otomatis deduct stok berdasarkan recipe',
    ]
    for item in create_recipe:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('C. Menu Management', style='Heading 3')
    menu_management = [
        'Tambah menu baru: Nama, kategori, harga, deskripsi, foto',
        'Edit menu existing',
        'Set menu availability (aktif/nonaktif)',
        'Jika stok bahan habis, sistem otomatis set menu unavailable',
    ]
    for item in menu_management:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.5. Riwayat Penjualan', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Sales History" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph('Melihat history semua transaksi penjualan dengan detail lengkap.')

    doc.add_paragraph('Fitur-fitur:', style='Heading 3')
    sales_history = [
        'Tabel semua transaksi',
        'Kolom: Tanggal, Order ID, Meja, Items, Total, Payment Method, Cashier',
        'Filter berdasarkan:',
        '   - Date range',
        '   - Payment method',
        '   - Cashier',
        'Search berdasarkan order ID',
        'Klik transaksi untuk lihat detail lengkap',
        'Export data ke Excel',
        'Cetak ulang struk jika diperlukan',
    ]
    for item in sales_history:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.6. Data Pelanggan', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Customer" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola database pelanggan untuk program loyalty, feedback, dan customer analytics.'
    )

    doc.add_paragraph('Modul Customer:', style='Heading 3')

    doc.add_paragraph('A. Customer Database', style='Heading 4')
    customer_db = [
        'Daftar semua pelanggan terdaftar',
        'Info: Nama, Email, Telepon, Total visits, Total spent',
        'Tambah customer baru',
        'Edit customer info',
        'Lihat purchase history per customer',
    ]
    for item in customer_db:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Membership/Loyalty Program', style='Heading 4')
    doc.add_paragraph('Akses: Customer → Membership')
    membership = [
        'Manage membership tiers (Bronze, Silver, Gold, Platinum)',
        'Set benefits per tier (discount, points, dll)',
        'Member card management',
        'Points accumulation rules',
        'Redeem points',
    ]
    for item in membership:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('C. Customer Feedback', style='Heading 4')
    doc.add_paragraph('Akses: Customer → Feedback')
    feedback = [
        'Daftar feedback/review dari pelanggan',
        'Rating: 1-5 bintang',
        'Kategori: Food, Service, Ambience, Price',
        'Comments',
        'Respond to feedback',
        'Analisis sentiment dan trends',
    ]
    for item in feedback:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.7. Jadwal Karyawan', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Schedule" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola jadwal shift karyawan (khususnya kasir) untuk validasi buka shift.'
    )

    doc.add_paragraph('Fitur-fitur:', style='Heading 3')
    schedule_features = [
        'Calendar view jadwal shift',
        'Assign shift ke karyawan:',
        '   - Pilih karyawan',
        '   - Pilih tanggal',
        '   - Pilih shift (Pagi/Siang/Sore/Malam)',
        '   - Jam mulai dan selesai',
        'Edit/Delete schedule',
        'View per karyawan atau per shift',
        'Export schedule ke PDF',
        'Sistem validasi saat kasir buka shift:',
        '   - Jika sesuai schedule: Auto approve',
        '   - Jika tidak sesuai: Perlu manager override',
    ]
    for item in schedule_features:
        if item.startswith('   -'):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        elif item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.8. Laporan', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Report" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph('Generate berbagai laporan untuk analisis bisnis dan decision making.')

    doc.add_paragraph('Jenis-jenis Laporan:', style='Heading 3')

    report_types = [
        ('Sales Report',
         'Laporan penjualan per periode dengan breakdown per menu, kategori, payment method, dan cashier.'),
        ('Daily Sales Summary',
         'Ringkasan penjualan harian: total sales, total orders, average order value, top selling items.'),
        ('Cashier Performance Report',
         'Performa kasir: total transaksi, total sales, average handling time, cash discrepancy.'),
        ('Shift Report',
         'Laporan per shift: durasi, total sales, total transaksi, cash reconciliation.'),
        ('Inventory Report',
         'Stock position, stock value, low stock items, expired items, stock movements.'),
        ('Menu Performance Report',
         'Analisis menu: best sellers, slow movers, contribution margin, popularity trends.'),
        ('Customer Analytics',
         'Customer behavior: repeat rate, average visit frequency, customer lifetime value.'),
        ('Financial Statement',
         'Income statement, revenue vs cost, profit margin.'),
    ]

    for report_name, description in report_types:
        p = doc.add_paragraph()
        p.add_run(f'{report_name}: ').bold = True
        p.add_run(description)

    doc.add_paragraph()
    doc.add_paragraph('Cara Generate Laporan:', style='Heading 3')
    generate_report = [
        'Pilih jenis laporan',
        'Pilih periode/date range',
        'Pilih filter tambahan (opsional): cashier, menu, payment method, dll',
        'Klik "Generate Report"',
        'Preview laporan',
        'Export:',
        '   - PDF (untuk print)',
        '   - Excel (untuk analisis)',
        '   - Send via Email',
    ]
    for i, item in enumerate(generate_report, 1):
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{item}', style='List Number')

    doc.add_heading('4.9. Pengaturan', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Office → Menu "Settings" di sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph('Konfigurasi pengaturan sistem restoran (hanya Admin).')

    doc.add_paragraph('Kategori Pengaturan:', style='Heading 3')

    settings_categories = [
        ('General Settings', [
            'Nama restoran',
            'Logo (upload)',
            'Alamat dan contact info',
            'Jam operasional',
            'Currency',
            'Timezone',
        ]),
        ('Tax Settings', [
            'Tax rate (%)',
            'Service charge (%)',
            'Apply to: Dine-in/Takeaway/Both',
        ]),
        ('Table Settings', [
            'Jumlah meja',
            'Nomor meja',
            'Kapasitas per meja',
        ]),
        ('Payment Settings', [
            'Metode pembayaran aktif',
            'Payment gateway config (QRIS, dll)',
        ]),
        ('Printer Settings', [
            'Printer untuk struk',
            'Printer dapur (kitchen print)',
            'Paper size',
        ]),
        ('User & Roles', [
            'Manage users',
            'Assign roles',
            'Permissions per role',
        ]),
    ]

    for category_name, settings_list in settings_categories:
        doc.add_paragraph(category_name, style='Heading 4')
        for setting in settings_list:
            doc.add_paragraph(setting, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Chapter 5: Kitchen
    doc.add_heading('5. DAPUR (KITCHEN)', 1)

    doc.add_heading('5.1. Dashboard Dapur', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph(
        'Setelah login sebagai staff dapur, Anda akan langsung masuk ke Dashboard Dapur.'
    )

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Dashboard untuk kitchen staff melihat pesanan baru yang masuk dan update status masakan.'
    )

    doc.add_paragraph('Tampilan Dashboard Dapur:', style='Heading 3')
    kitchen_dashboard = [
        'Layout berbasis kartu (card) untuk setiap pesanan',
        'Badge status berwarna:',
        '   - Biru: CONFIRMED (pesanan baru)',
        '   - Oranye: PREPARING (sedang dimasak)',
        '   - Hijau: READY (siap diantar)',
        'Informasi per kartu:',
        '   - Nomor pesanan',
        '   - Nomor meja atau Takeaway',
        '   - Daftar item dan quantity',
        '   - Catatan khusus (jika ada)',
        '   - Durasi pesanan (elapsed time)',
        'Sound notification saat pesanan baru masuk (opsional)',
        'Auto-refresh untuk update real-time',
    ]
    for item in kitchen_dashboard:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2. Menerima Pesanan', 2)

    doc.add_paragraph('Workflow menerima pesanan baru:')
    receive_order = [
        'Pesanan baru masuk dengan status CONFIRMED (badge biru)',
        'Kitchen staff lihat detail pesanan dan catatan',
        'Klik "Mulai Memasak" atau "Start Preparing"',
        'Status berubah menjadi PREPARING (badge oranye)',
        'Timer mulai berjalan untuk tracking cooking time',
        'Kartu berpindah ke section "In Progress"',
    ]
    for step in receive_order:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('5.3. Update Status Pesanan', 2)

    doc.add_paragraph('Setelah masakan selesai:')
    update_status = [
        'Klik "Siap" atau "Mark as Ready" di kartu pesanan',
        'Status berubah menjadi READY (badge hijau)',
        'Kartu berpindah ke section "Ready to Serve"',
        'Sistem catat finish time dan cooking duration',
        'Waiter/kasir akan melihat pesanan siap diantar',
        'Waiter mengambil makanan dan antar ke meja',
        'Waiter/kasir update status menjadi COMPLETED (sudah diantar)',
        'Kartu pesanan hilang dari dashboard dapur',
    ]
    for step in update_status:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Tips untuk Kitchen Staff:', style='Heading 3')
    kitchen_tips = [
        'Prioritaskan pesanan berdasarkan durasi (merah = lama menunggu)',
        'Baca catatan khusus dengan teliti (tidak pedas, alergi, dll)',
        'Update status secara real-time untuk koordinasi dengan front',
        'Jika ada masalah (bahan habis, dll), inform kasir segera',
        'Check dashboard secara berkala untuk pesanan baru',
    ]
    for tip in kitchen_tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_page_break()

    # Chapter 6: Common Features
    doc.add_heading('6. FITUR UMUM', 1)

    doc.add_heading('6.1. Profil Pengguna', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik menu hamburger (☰) di sidebar → Pilih "Profile"')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph('Melihat dan mengedit informasi profil Anda.')

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    profile_info = [
        'Personal Information:',
        '   - Nama lengkap',
        '   - Email',
        '   - Nomor telepon',
        'Employment Information:',
        '   - Role (Admin/Manager/Cashier/Kitchen/Warehouse)',
        '   - Branch',
        'Account Information:',
        '   - Email login',
        '   - Last login time',
    ]
    for item in profile_info:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('Mengedit Profil:', style='Heading 3')
    edit_profile = [
        'Klik "Edit Profile"',
        'Update informasi yang bisa diubah (nama, telepon)',
        'Klik "Save Changes"',
        'Catatan: Email dan role hanya bisa diubah oleh Admin',
    ]
    for step in edit_profile:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Mengubah Password:', style='Heading 3')
    change_password = [
        'Klik "Change Password"',
        'Masukkan password lama',
        'Masukkan password baru',
        'Konfirmasi password baru',
        'Klik "Update Password"',
        'Password requirements: minimal 8 karakter',
    ]
    for step in change_password:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('6.2. Pengaturan (User Settings)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik menu hamburger (☰) di sidebar → Pilih "Settings"')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph('Pengaturan personal user (berbeda dengan Settings di Office yang untuk sistem).')

    doc.add_paragraph('Opsi yang tersedia:', style='Heading 3')
    user_settings = [
        'Notifications: Enable/disable notifikasi',
        'Sound: Enable/disable sound untuk notifikasi',
        'Language: Bahasa Indonesia (default)',
        'Theme: Light/Dark mode (jika ada)',
        'Print Settings: Pilih printer default',
        'Auto-logout: Set waktu auto-logout saat idle',
    ]
    for item in user_settings:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3. Logout', 2)

    doc.add_paragraph('Cara logout:', style='Heading 3')
    logout_steps = [
        'Klik menu hamburger (☰) di sidebar',
        'Pilih "Logout" (paling bawah)',
        'Konfirmasi jika diminta',
        'Anda akan diarahkan kembali ke halaman login',
        'Session dihapus dari sistem',
    ]
    for step in logout_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('PENTING:', style='Heading 3')
    logout_notes = [
        'Jika Anda seorang Kasir dengan shift aktif:',
        '   - Sistem akan warning bahwa Anda punya shift aktif',
        '   - Anda masih bisa logout, tapi shift tetap aktif',
        '   - Pastikan tutup shift terlebih dahulu sebelum logout (best practice)',
        'Selalu logout setelah selesai bekerja untuk keamanan',
        'Jangan share kredensial login dengan orang lain',
    ]
    for note in logout_notes:
        if note.startswith('   '):
            doc.add_paragraph(note.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(note, style='List Bullet')

    doc.add_page_break()

    # Chapter 7: Tips & Troubleshooting
    doc.add_heading('7. TIPS & TROUBLESHOOTING', 1)

    doc.add_heading('Tips Penggunaan Sistem', 2)

    tips = [
        ('Selalu Buka Shift Terlebih Dahulu',
         'Kasir HARUS buka shift sebelum terima pesanan dan proses pembayaran. Tanpa shift aktif, '
         'transaksi tidak akan tercatat dengan benar.'),
        ('Pastikan Status Pesanan Benar',
         'Hanya pesanan COMPLETED yang bisa dibayar. Koordinasi dengan dapur untuk update status.'),
        ('Hitung Modal Awal dengan Teliti',
         'Modal awal di awal shift akan dibandingkan dengan actual cash saat tutup shift. '
         'Selisih besar akan memicu audit.'),
        ('Check Stok Sebelum Terima Pesanan',
         'Jika bahan habis, menu akan otomatis unavailable. Inform pelanggan sejak awal.'),
        ('Gunakan Catatan untuk Permintaan Khusus',
         'Manfaatkan field catatan saat buat pesanan untuk komunikasi dengan dapur.'),
        ('Print Struk untuk Pelanggan',
         'Selalu berikan struk pembayaran ke pelanggan sebagai bukti transaksi.'),
        ('Tutup Shift Setiap Akhir Jam Kerja',
         'Jangan lupa tutup shift untuk reconciliation dan generate laporan.'),
        ('Backup Data Penting',
         'Export laporan penting ke Excel/PDF secara berkala.'),
    ]

    for tip_title, tip_desc in tips:
        p = doc.add_paragraph()
        p.add_run(f'{tip_title}: ').bold = True
        p.add_run(tip_desc)

    doc.add_heading('Troubleshooting - Masalah Umum', 2)

    doc.add_paragraph('Masalah 1: Tidak Bisa Login', style='Heading 3')
    login_troubleshoot = [
        'Periksa: Email dan password benar? (case-sensitive)',
        'Periksa: Caps Lock tidak aktif?',
        'Periksa: Koneksi internet stabil?',
        'Coba: Clear browser cache dan cookies',
        'Coba: Gunakan browser lain (Chrome, Firefox, Safari)',
        'Jika masih gagal: Hubungi administrator',
    ]
    for item in login_troubleshoot:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 2: Tidak Bisa Buka Shift', style='Heading 3')
    shift_troubleshoot = [
        'Kemungkinan: Anda tidak ada di jadwal shift hari ini',
        'Solusi 1: Hubungi Manager untuk approval (override)',
        'Solusi 2: Minta Admin update jadwal Anda',
        'Kemungkinan: Anda masih punya shift aktif yang belum ditutup',
        'Solusi: Tutup shift lama terlebih dahulu',
    ]
    for item in shift_troubleshoot:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 3: Pesanan Tidak Muncul di Transaksi', style='Heading 3')
    order_not_appear = [
        'Kemungkinan: Status pesanan belum COMPLETED',
        'Cek: Status pesanan di halaman Meja atau tanya dapur',
        'Solusi: Update status pesanan menjadi COMPLETED terlebih dahulu',
        'Kemungkinan: Pesanan sudah dibayar',
        'Cek: Riwayat penjualan untuk confirm',
    ]
    for item in order_not_appear:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 4: Tidak Bisa Tutup Shift', style='Heading 3')
    close_shift_issue = [
        'Kemungkinan: Masih ada pesanan unpaid/unsettled',
        'Solusi: Check semua pesanan sudah dibayar (status PAID)',
        'Solusi: Proses pembayaran untuk pesanan yang belum bayar',
        'Kemungkinan: Ada transaksi yang pending',
        'Solusi: Selesaikan atau batalkan transaksi pending',
    ]
    for item in close_shift_issue:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 5: Struk Tidak Mau Print', style='Heading 3')
    print_issue = [
        'Periksa: Printer menyala dan terhubung',
        'Periksa: Ada kertas di printer',
        'Periksa: Printer driver ter-install',
        'Coba: Print test page dari settings printer',
        'Coba: Gunakan browser lain',
        'Solusi alternatif: Save struk as PDF dan print manual',
    ]
    for item in print_issue:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 6: Menu Tidak Tersedia (Unavailable)', style='Heading 3')
    menu_unavailable = [
        'Kemungkinan: Stok bahan baku habis',
        'Cek: Stok management untuk confirm',
        'Solusi: Stock in bahan yang dibutuhkan',
        'Solusi alternatif: Set menu available manual (jika stok cukup tapi sistem error)',
        'Kemungkinan: Menu memang diset unavailable oleh admin',
        'Solusi: Hubungi admin untuk aktivasi',
    ]
    for item in menu_unavailable:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Browser yang Direkomendasikan', 2)
    browsers = [
        'Google Chrome (versi 90+) - Recommended',
        'Mozilla Firefox (versi 88+)',
        'Safari (versi 14+) untuk Mac',
        'Microsoft Edge (versi 90+)',
    ]
    doc.add_paragraph('Sistem dioptimalkan untuk browser modern. Gunakan salah satu browser berikut:')
    for browser in browsers:
        doc.add_paragraph(browser, style='List Bullet')

    doc.add_heading('Kontak Bantuan', 2)

    contact_help = [
        ('IT Support', 'support@ladapala-resto.com | Ext: 101'),
        ('Admin/Manager', 'admin@ladapala-resto.com | Ext: 102'),
    ]

    doc.add_paragraph('Jika mengalami masalah atau butuh bantuan:')
    for contact_title, contact_info in contact_help:
        p = doc.add_paragraph()
        p.add_run(f'{contact_title}: ').bold = True
        p.add_run(contact_info)

    doc.add_paragraph()
    doc.add_paragraph(
        'Saat menghubungi support, siapkan informasi: Nama, Role, Deskripsi masalah, '
        'Screenshot (jika ada), dan langkah-langkah yang sudah dicoba.'
    )

    doc.add_page_break()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer.add_run(
        '\n\n───────────────────────────────────────\n'
        'PANDUAN PENGGUNAAN SISTEM POS RESTORAN LADAPALA\n'
        'Versi 1.0 - November 2025\n'
        '© 2025 Ladapala Restaurant. All rights reserved.\n'
        '───────────────────────────────────────'
    )
    footer_text.font.size = Pt(9)
    footer_text.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = '/Users/pro/Dev/ladapala/Panduan_Sistem_POS_Restoran_Ladapala.docx'
    doc.save(output_path)
    print(f'✓ Manual restoran berhasil dibuat: {output_path}')
    return output_path

if __name__ == '__main__':
    create_manual()
