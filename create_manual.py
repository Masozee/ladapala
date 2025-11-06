#!/usr/bin/env python3
"""
Script to create comprehensive user manual for Kapulaga Hotel Management System
in Bahasa Indonesia
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Set document title
    title = doc.add_heading('PANDUAN PENGGUNAAN SISTEM MANAJEMEN HOTEL KAPULAGA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version info
    version = doc.add_paragraph('Versi 1.0 - November 2025')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('DAFTAR ISI', 1)
    toc_items = [
        '1. Pendahuluan',
        '2. Memulai Sistem',
        '   2.1. Login ke Sistem',
        '   2.2. Mengenal Dashboard',
        '3. Front Desk - Operasional Harian',
        '   3.1. Dashboard Front Desk',
        '   3.2. Manajemen Booking',
        '   3.3. Status Kamar',
        '   3.4. Keluhan Tamu',
        '   3.5. Pembayaran',
        '4. Office - Manajemen & Administrasi',
        '   4.1. Office Dashboard',
        '   4.2. Database Tamu',
        '   4.3. Manajemen Karyawan',
        '   4.4. Housekeeping',
        '   4.5. Keuangan',
        '   4.6. Gudang',
        '   4.7. Supplier',
        '   4.8. Laporan',
        '   4.9. Pengaturan Sistem',
        '5. Support - Housekeeping & Maintenance',
        '   5.1. Support Dashboard',
        '   5.2. Maintenance',
        '   5.3. Housekeeping',
        '   5.4. Permintaan Amenities',
        '   5.5. Laporan Support',
        '6. Fitur Umum',
        '   6.1. Profil Pengguna',
        '   6.2. Kalender',
        '   6.3. Logout',
        '7. Tips & Troubleshooting',
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('1. PENDAHULUAN', 1)

    doc.add_heading('Tentang Sistem', 2)
    doc.add_paragraph(
        'Sistem Manajemen Hotel Kapulaga adalah platform terintegrasi yang dirancang khusus '
        'untuk memudahkan pengelolaan operasional hotel. Sistem ini menggabungkan seluruh '
        'aspek operasional hotel mulai dari reservasi, manajemen kamar, housekeeping, '
        'maintenance, keuangan, hingga pelaporan dalam satu platform yang mudah digunakan.'
    )

    doc.add_heading('Tiga Area Utama Sistem', 2)

    areas = [
        ('Front Desk (Hijau Tosca)',
         'Area untuk operasional harian front desk seperti check-in, check-out, booking, '
         'dan melayani tamu. Warna hijau tosca (#005357) menandakan area ini.'),
        ('Office (Biru)',
         'Area untuk manajemen dan administrasi seperti database tamu, karyawan, keuangan, '
         'gudang, dan pengaturan sistem. Warna biru (#4E61D3) menandakan area ini.'),
        ('Support (Oranye)',
         'Area untuk tim support seperti housekeeping, maintenance, dan permintaan amenities. '
         'Warna oranye (#F87B1B) menandakan area ini.'),
    ]

    for area_name, description in areas:
        p = doc.add_paragraph()
        p.add_run(area_name).bold = True
        p.add_run(f'\n{description}')

    doc.add_heading('Hak Akses Berdasarkan Departemen', 2)
    doc.add_paragraph(
        'Setiap pengguna memiliki hak akses berbeda berdasarkan departemen mereka:'
    )

    access_list = [
        'Management/Admin: Akses penuh ke semua area (Front Desk, Office, dan Support)',
        'Front Office/Reception: Akses ke area Front Desk',
        'Housekeeping: Akses ke area Support',
        'Maintenance: Akses ke area Support',
        'Finance: Akses ke area Office',
    ]

    for access in access_list:
        doc.add_paragraph(access, style='List Bullet')

    doc.add_page_break()

    # Chapter 2: Getting Started
    doc.add_heading('2. MEMULAI SISTEM', 1)

    doc.add_heading('2.1. Login ke Sistem', 2)

    doc.add_paragraph('Langkah-langkah untuk login:')

    login_steps = [
        'Buka browser (Chrome, Firefox, Safari, atau Edge)',
        'Ketik alamat: http://localhost:3000/login',
        'Anda akan melihat halaman login dengan dua kolom:',
        '   - Kolom kiri: Form login',
        '   - Kolom kanan: Informasi hotel dan kontak bantuan',
        'Masukkan email Anda di kolom "Email"',
        'Masukkan password Anda di kolom "Password"',
        'Klik tombol "Sign in" berwarna hijau tosca',
        'Jika berhasil, Anda akan diarahkan ke dashboard sesuai hak akses Anda',
    ]

    for i, step in enumerate(login_steps, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Catatan Penting:', style='Heading 3')
    notes = [
        'Pastikan Anda menggunakan email dan password yang telah diberikan oleh administrator',
        'Jika lupa password, hubungi administrator melalui kontak yang tertera di halaman login',
        'Sistem akan mengingat sesi Anda, jadi Anda tidak perlu login ulang setiap kali membuka browser',
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_heading('2.2. Mengenal Dashboard', 2)

    doc.add_paragraph(
        'Setelah login berhasil, Anda akan melihat dashboard utama. Dashboard yang muncul '
        'tergantung pada departemen dan hak akses Anda:'
    )

    dashboard_types = [
        ('Front Desk Dashboard', 'Untuk staff front office/reception'),
        ('Office Dashboard', 'Untuk management dan staff administrasi'),
        ('Support Dashboard', 'Untuk staff housekeeping dan maintenance'),
    ]

    for dash_name, dash_desc in dashboard_types:
        p = doc.add_paragraph()
        p.add_run(f'{dash_name}: ').bold = True
        p.add_run(dash_desc)

    doc.add_heading('Elemen-elemen Dashboard', 3)

    doc.add_paragraph('1. Sidebar (Kolom Kiri):')
    sidebar_items = [
        'Logo hotel di bagian atas',
        'Ikon-ikon menu navigasi (hover untuk melihat nama menu)',
        'Badge merah menunjukkan jumlah item yang perlu perhatian',
        'Menu Kalender dan Profile di bagian bawah',
    ]
    for item in sidebar_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('2. Top Navbar (Bagian Atas):')
    navbar_items = [
        'Breadcrumb: Menunjukkan lokasi Anda saat ini di sistem',
        'Tanggal dan Waktu: Real-time di pojok kanan',
        'Ikon Pencarian: Untuk mencari data dengan cepat',
        'Ikon Kalender: Melihat event hari ini',
        'Ikon Notifikasi: Menampilkan notifikasi (badge merah menunjukkan jumlah)',
        'Ikon Dark Mode: Mengubah tema gelap/terang',
        'Ikon Logout (merah): Untuk keluar dari sistem',
    ]
    for item in navbar_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('3. Konten Utama (Tengah):')
    doc.add_paragraph(
        'Area ini menampilkan informasi utama seperti statistik, tabel data, '
        'dan form input sesuai dengan menu yang Anda pilih.'
    )

    doc.add_page_break()

    # Chapter 3: Front Desk Operations
    doc.add_heading('3. FRONT DESK - OPERASIONAL HARIAN', 1)

    doc.add_heading('3.1. Dashboard Front Desk', 2)
    doc.add_paragraph(
        'Dashboard Front Desk adalah halaman utama untuk staff reception. '
        'Di sini Anda dapat melihat ringkasan operasional hotel hari ini.'
    )

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Dashboard" (rumah) di sidebar kiri')

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    dashboard_info = [
        'Total Kamar Terisi: Jumlah kamar yang saat ini ditempati tamu',
        'Total Booking Hari Ini: Jumlah reservasi yang dijadwalkan hari ini',
        'Pending Bookings: Booking yang menunggu konfirmasi',
        'Keluhan Belum Selesai: Jumlah komplain tamu yang masih dalam proses',
        'Grafik Okupansi: Visualisasi tingkat hunian kamar',
        'Daftar Check-in Hari Ini: Tamu yang akan check-in',
        'Daftar Check-out Hari Ini: Tamu yang akan check-out',
    ]
    for info in dashboard_info:
        doc.add_paragraph(info, style='List Bullet')

    doc.add_heading('3.2. Manajemen Booking', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Bookings" (kalender) di sidebar kiri')

    doc.add_paragraph('A. Melihat Daftar Booking', style='Heading 3')
    booking_view = [
        'Halaman akan menampilkan tabel semua booking',
        'Gunakan filter status untuk menyaring: Pending, Confirmed, Checked In, Checked Out, Cancelled',
        'Gunakan search bar untuk mencari berdasarkan nama tamu, nomor reservasi, atau kamar',
        'Badge status berwarna menunjukkan status booking:',
        '   - Kuning: Pending (menunggu konfirmasi)',
        '   - Biru: Confirmed (sudah dikonfirmasi)',
        '   - Hijau: Checked In (tamu sudah check-in)',
        '   - Abu-abu: Checked Out (tamu sudah check-out)',
        '   - Merah: Cancelled (dibatalkan)',
    ]
    for i, item in enumerate(booking_view, 1):
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{item}')

    doc.add_paragraph()
    doc.add_paragraph('B. Membuat Booking Baru', style='Heading 3')
    doc.add_paragraph('Langkah-langkah:')
    new_booking = [
        'Klik tombol "+ New Booking" di pojok kanan atas',
        'Isi form booking:',
        '   - Data Tamu: Nama, email, telepon (atau pilih dari database tamu)',
        '   - Tanggal Check-in dan Check-out',
        '   - Tipe Kamar: Pilih dari dropdown',
        '   - Jumlah Tamu',
        '   - Special Requests: Permintaan khusus (opsional)',
        'Sistem akan menampilkan harga total otomatis',
        'Klik "Create Booking" untuk menyimpan',
        'Booking akan tersimpan dengan status "Pending"',
    ]
    for i, step in enumerate(new_booking, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Konfirmasi Booking', style='Heading 3')
    confirm_booking = [
        'Klik ikon titik tiga (⋯) di baris booking yang akan dikonfirmasi',
        'Pilih "Confirm Booking"',
        'Status akan berubah menjadi "Confirmed" (biru)',
        'Sistem akan mengirim notifikasi konfirmasi ke email tamu (jika dikonfigurasi)',
    ]
    for step in confirm_booking:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('D. Check-in Tamu', style='Heading 3')
    checkin_steps = [
        'Pastikan booking sudah berstatus "Confirmed"',
        'Klik ikon titik tiga (⋯) di baris booking',
        'Pilih "Check In"',
        'Verifikasi data tamu',
        'Pilih nomor kamar yang akan ditempati',
        'Klik "Check In" untuk konfirmasi',
        'Status booking berubah menjadi "Checked In" (hijau)',
        'Kamar akan otomatis berubah status menjadi "Occupied"',
    ]
    for step in checkin_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('E. Check-out Tamu', style='Heading 3')
    checkout_steps = [
        'Klik ikon titik tiga (⋯) di baris booking yang akan check-out',
        'Pilih "Check Out"',
        'Sistem akan menampilkan ringkasan tagihan',
        'Pastikan semua pembayaran sudah lunas',
        'Klik "Check Out" untuk konfirmasi',
        'Status booking berubah menjadi "Checked Out" (abu-abu)',
        'Kamar akan otomatis berubah status menjadi "Dirty" (perlu dibersihkan)',
    ]
    for step in checkout_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('F. Membatalkan Booking', style='Heading 3')
    cancel_steps = [
        'Klik ikon titik tiga (⋯) di baris booking yang akan dibatalkan',
        'Pilih "Cancel Booking"',
        'Masukkan alasan pembatalan (opsional)',
        'Klik "Confirm Cancellation"',
        'Status booking berubah menjadi "Cancelled" (merah)',
    ]
    for step in cancel_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    doc.add_heading('3.3. Status Kamar (Room Status)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Room Status" (kasur) di sidebar kiri')

    doc.add_paragraph('Fungsi halaman ini:', style='Heading 3')
    doc.add_paragraph(
        'Memonitor status semua kamar hotel secara real-time untuk memudahkan '
        'koordinasi dengan housekeeping dan maintenance.'
    )

    doc.add_paragraph('Status Kamar yang Ada:', style='Heading 3')
    room_statuses = [
        'Available (Hijau): Kamar kosong dan siap ditempati',
        'Occupied (Biru): Kamar sedang ditempati tamu',
        'Dirty (Kuning): Kamar kotor, perlu dibersihkan',
        'Cleaning (Oranye): Kamar sedang dibersihkan housekeeping',
        'Maintenance (Merah): Kamar sedang dalam perbaikan',
        'Reserved (Ungu): Kamar sudah dibooking tapi belum check-in',
    ]
    for status in room_statuses:
        doc.add_paragraph(status, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Fitur-fitur:', style='Heading 3')
    room_features = [
        'Filter berdasarkan tipe kamar (Standard, Deluxe, Suite, dll)',
        'Filter berdasarkan status kamar',
        'Search kamar berdasarkan nomor kamar',
        'View mode: Grid (kartu) atau List (tabel)',
        'Klik kartu kamar untuk melihat detail lengkap',
        'Update status kamar langsung dari halaman ini',
    ]
    for feature in room_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Mengubah Status Kamar:', style='Heading 3')
    change_status = [
        'Klik pada kartu kamar yang ingin diubah',
        'Modal detail kamar akan muncul',
        'Pilih status baru dari dropdown',
        'Klik "Update Status"',
        'Perubahan akan langsung terlihat di dashboard',
    ]
    for step in change_status:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('3.4. Keluhan Tamu (Complaints)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Complaints" (tanda tanya) di sidebar kiri')

    doc.add_paragraph('Fungsi halaman ini:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola dan melacak semua keluhan atau permintaan dari tamu hotel '
        'untuk memastikan pelayanan yang optimal.'
    )

    doc.add_paragraph('A. Melihat Daftar Keluhan', style='Heading 3')
    complaints_view = [
        'Tabel menampilkan semua keluhan tamu',
        'Informasi yang ditampilkan: Nomor complaint, Nama tamu, Kamar, Kategori, Status, Tanggal',
        'Filter berdasarkan status: Pending, In Progress, Resolved',
        'Filter berdasarkan kategori: Room, Service, Facility, Food, Other',
        'Search berdasarkan nama tamu atau nomor complaint',
    ]
    for item in complaints_view:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Membuat Keluhan Baru', style='Heading 3')
    new_complaint = [
        'Klik tombol "+ New Complaint"',
        'Isi form:',
        '   - Pilih Tamu: Dari dropdown tamu yang sedang check-in',
        '   - Kategori: Room/Service/Facility/Food/Other',
        '   - Deskripsi: Jelaskan detail keluhan',
        '   - Prioritas: Low/Medium/High',
        'Klik "Submit Complaint"',
        'Keluhan tersimpan dengan status "Pending"',
        'Tim terkait akan menerima notifikasi',
    ]
    for i, step in enumerate(new_complaint, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Menindaklanjuti Keluhan', style='Heading 3')
    followup_complaint = [
        'Klik pada baris keluhan untuk melihat detail',
        'Update status menjadi "In Progress" saat mulai ditangani',
        'Tambahkan catatan tindakan yang dilakukan',
        'Setelah selesai, ubah status menjadi "Resolved"',
        'Isi form resolusi dengan penjelasan penyelesaian',
        'Sistem akan mencatat waktu penyelesaian otomatis',
    ]
    for step in followup_complaint:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('3.5. Pembayaran (Payments)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Payments" (kartu kredit) di sidebar kiri')

    doc.add_paragraph('Fungsi halaman ini:', style='Heading 3')
    doc.add_paragraph(
        'Memproses dan mencatat semua transaksi pembayaran dari tamu, baik untuk '
        'booking, layanan tambahan, maupun minibar.'
    )

    doc.add_paragraph('A. Memproses Pembayaran', style='Heading 3')
    payment_process = [
        'Pilih booking yang akan dibayar dari daftar',
        'Sistem menampilkan rincian tagihan:',
        '   - Biaya kamar (sesuai durasi menginap)',
        '   - Layanan tambahan (jika ada)',
        '   - Tax dan service charge',
        '   - Total yang harus dibayar',
        'Pilih metode pembayaran:',
        '   - Cash (Tunai)',
        '   - Credit Card (Kartu Kredit)',
        '   - Debit Card (Kartu Debit)',
        '   - Bank Transfer',
        'Masukkan jumlah yang dibayarkan',
        'Sistem akan hitung kembalian otomatis (untuk cash)',
        'Klik "Process Payment"',
        'Cetak atau kirim struk pembayaran ke email tamu',
    ]
    for i, step in enumerate(payment_process, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('B. Voucher dan Diskon', style='Heading 3')
    doc.add_paragraph(
        'Sistem mendukung pemberian voucher dan diskon otomatis untuk meningkatkan '
        'kepuasan tamu dan program loyalitas.'
    )

    doc.add_paragraph('Jenis Diskon:', style='Heading 4')
    discount_types = [
        'Voucher Code: Tamu masukkan kode voucher (contoh: PROMO2025, PROMO50)',
        'Diskon Otomatis: Sistem otomatis memberikan diskon berdasarkan:',
        '   - Early Bird: Booking jauh hari',
        '   - Length of Stay: Menginap lebih dari X malam',
        '   - Seasonal Discount: Promo musiman',
        'Loyalty Points: Tamu redeem poin loyalitas (1 poin = Rp 1)',
    ]
    for item in discount_types:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Cara Menggunakan Voucher:', style='Heading 4')
    voucher_steps = [
        'Saat proses pembayaran, masukkan kode voucher di kolom "Voucher Code"',
        'Klik "Apply" atau tekan Enter',
        'Sistem akan validasi voucher:',
        '   - Cek apakah kode valid',
        '   - Cek masa berlaku',
        '   - Cek apakah sudah digunakan (jika single-use)',
        '   - Cek minimum booking requirements',
        'Jika valid, diskon akan otomatis diterapkan',
        'Rincian diskon akan muncul di breakdown pembayaran',
        'Total yang harus dibayar akan berkurang',
    ]
    for i, step in enumerate(voucher_steps, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Perhitungan Pembayaran dengan Diskon:', style='Heading 4')
    doc.add_paragraph(
        'Sistem secara otomatis menghitung status pembayaran dengan mempertimbangkan '
        'semua diskon yang diterapkan:'
    )
    payment_calc = [
        'Grand Total: Total kamar + Tax (21%) + Service Charge (10%)',
        'Dikurangi: Voucher Discount',
        'Dikurangi: Automatic Discount',
        'Dikurangi: Loyalty Points Value',
        'Hasil: Final Amount (jumlah yang harus dibayar tamu)',
        'Status "Fully Paid" akan muncul jika total bayaran ≥ Final Amount',
    ]
    for item in payment_calc:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Catatan Penting: ').bold = True
    p.add_run(
        'Update sistem November 2025 memastikan status "Fully Paid" sudah memperhitungkan '
        'semua voucher dan diskon. Tamu yang membayar dengan voucher akan langsung '
        'ditandai sebagai fully paid jika pembayaran sudah memenuhi Final Amount.'
    )

    doc.add_paragraph()
    doc.add_paragraph('C. Melihat Riwayat Pembayaran', style='Heading 3')
    payment_history = [
        'Tab "Payment History" menampilkan semua transaksi',
        'Filter berdasarkan tanggal',
        'Filter berdasarkan metode pembayaran',
        'Export data ke Excel atau PDF untuk laporan',
        'Cetak ulang struk pembayaran jika diperlukan',
    ]
    for item in payment_history:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Chapter 4: Office Operations
    doc.add_heading('4. OFFICE - MANAJEMEN & ADMINISTRASI', 1)

    doc.add_heading('4.1. Office Dashboard', 2)
    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Office" (gedung) di sidebar Front Desk, atau klik ikon "Office Dashboard" di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Dashboard Office adalah pusat kontrol untuk manajemen dan administrasi hotel. '
        'Menampilkan metrics penting, grafik performa, dan akses cepat ke semua modul administrasi.'
    )

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    office_dashboard = [
        'Total Revenue bulan ini',
        'Occupancy Rate (tingkat hunian)',
        'Total Employees',
        'Pending Tasks',
        'Grafik Revenue Trends',
        'Grafik Occupancy Trends',
        'Top Performing Room Types',
        'Housekeeping Performance',
        'Low Stock Alerts',
    ]
    for item in office_dashboard:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.2. Database Tamu (Guest Database)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Guest Database" (user multiple) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Menyimpan dan mengelola data semua tamu yang pernah menginap di hotel. '
        'Database ini berguna untuk marketing, loyalty program, dan analisis pelanggan.'
    )

    doc.add_paragraph('A. Melihat Data Tamu', style='Heading 3')
    guest_view = [
        'Tabel menampilkan semua tamu yang terdaftar',
        'Kolom: Nama, Email, Telepon, Nationality, Total Bookings, Last Visit',
        'Search berdasarkan nama, email, atau telepon',
        'Filter berdasarkan nationality',
        'Sort berdasarkan total bookings atau last visit',
    ]
    for item in guest_view:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Menambah Tamu Baru', style='Heading 3')
    new_guest = [
        'Klik tombol "+ Add Guest"',
        'Isi form data tamu:',
        '   - Personal Info: Nama, Email, Telepon',
        '   - Identification: ID Type (KTP/Passport), ID Number',
        '   - Address: Alamat lengkap, Kota, Negara',
        '   - Preferences: Tipe kamar favorit, special requests',
        'Klik "Save Guest"',
        'Tamu baru akan muncul di database',
    ]
    for i, step in enumerate(new_guest, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Edit Data Tamu', style='Heading 3')
    edit_guest = [
        'Klik pada nama tamu untuk melihat detail',
        'Klik tombol "Edit" di pojok kanan atas',
        'Update data yang perlu diubah',
        'Klik "Save Changes"',
    ]
    for step in edit_guest:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('D. Melihat Riwayat Booking Tamu', style='Heading 3')
    doc.add_paragraph(
        'Di halaman detail tamu, bagian bawah menampilkan semua riwayat booking tamu tersebut '
        'termasuk tanggal menginap, tipe kamar, dan total pembayaran.'
    )

    doc.add_heading('4.3. Manajemen Karyawan (Employees)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Employees" (user settings) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola data karyawan hotel termasuk informasi personal, jabatan, gaji, '
        'cuti, dan performa kerja.'
    )

    doc.add_paragraph('A. Melihat Daftar Karyawan', style='Heading 3')
    employee_view = [
        'Tabel menampilkan semua karyawan',
        'Kolom: Photo, Nama, Jabatan, Departemen, Status, Email, Telepon',
        'Filter berdasarkan departemen: Front Office, Housekeeping, F&B, Maintenance, Management',
        'Filter berdasarkan status: Active, On Leave, Terminated',
        'Search berdasarkan nama atau employee ID',
    ]
    for item in employee_view:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Menambah Karyawan Baru', style='Heading 3')
    new_employee = [
        'Klik tombol "+ Add Employee"',
        'Isi form data karyawan:',
        '   Tab Personal Info:',
        '     - Nama lengkap, Email, Telepon',
        '     - Tanggal lahir, Gender',
        '     - Alamat lengkap',
        '     - Upload foto (opsional)',
        '   Tab Employment Info:',
        '     - Employee ID',
        '     - Departemen',
        '     - Jabatan/Position',
        '     - Tanggal bergabung',
        '     - Gaji/Salary',
        '     - Employment Type: Full-time/Part-time/Contract',
        '   Tab Account Access:',
        '     - Email untuk login',
        '     - Password sementara',
        '     - Hak akses berdasarkan departemen',
        'Klik "Create Employee"',
        'Karyawan baru akan menerima email berisi kredensial login',
    ]
    for i, step in enumerate(new_employee, 1):
        if step.startswith('     -'):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        elif step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Edit Data Karyawan', style='Heading 3')
    edit_employee = [
        'Klik pada nama karyawan',
        'Klik tombol "Edit"',
        'Update data yang perlu diubah',
        'Klik "Save Changes"',
    ]
    for step in edit_employee:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('D. Menonaktifkan Karyawan', style='Heading 3')
    deactivate_employee = [
        'Klik pada nama karyawan',
        'Klik tombol titik tiga (⋯)',
        'Pilih "Deactivate" atau "Terminate"',
        'Masukkan alasan dan tanggal efektif',
        'Klik "Confirm"',
        'Akun karyawan akan dinonaktifkan',
    ]
    for step in deactivate_employee:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('4.4. Housekeeping (Office View)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Housekeeping" (circular reload) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Memonitor dan mengelola tugas housekeeping dari perspektif manajemen. '
        'Berbeda dengan Support Housekeeping yang digunakan oleh staff housekeeping, '
        'ini adalah view untuk supervisor dan manager.'
    )

    doc.add_paragraph('Fitur-fitur:', style='Heading 3')
    office_housekeeping = [
        'Dashboard housekeeping dengan statistik:',
        '   - Total tugas hari ini',
        '   - Tugas selesai',
        '   - Tugas dalam proses',
        '   - Tugas tertunda',
        'Assign tugas ke staff housekeeping',
        'Monitor performa staff (waktu penyelesaian, rating)',
        'Lihat history cleaning semua kamar',
        'Generate laporan housekeeping',
    ]
    for i, item in enumerate(office_housekeeping):
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.5. Keuangan (Financial)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Financial" (credit card) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola semua aspek keuangan hotel termasuk revenue, expenses, invoicing, '
        'dan financial reports.'
    )

    doc.add_paragraph('Modul-modul Financial:', style='Heading 3')

    doc.add_paragraph('A. Revenue Management', style='Heading 4')
    revenue_mgmt = [
        'Dashboard revenue dengan metrics:',
        '   - Total revenue hari ini/bulan ini/tahun ini',
        '   - Revenue by source: Room, F&B, Services',
        '   - Average Daily Rate (ADR)',
        '   - Revenue Per Available Room (RevPAR)',
        'Grafik revenue trends',
        'Revenue forecast berdasarkan booking',
        'Comparison dengan periode sebelumnya',
    ]
    for item in revenue_mgmt:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Expense Tracking', style='Heading 4')
    expense_tracking = [
        'Catat semua pengeluaran hotel',
        'Kategori expense: Payroll, Utilities, Supplies, Maintenance, Marketing, Other',
        'Upload bukti pembayaran (foto/scan)',
        'Approval workflow untuk expense besar',
        'Monthly expense report',
    ]
    for item in expense_tracking:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('C. Invoicing', style='Heading 4')
    invoicing = [
        'Generate invoice untuk corporate clients',
        'Invoice template yang customizable',
        'Track invoice status: Draft, Sent, Paid, Overdue',
        'Send invoice via email otomatis',
        'Record payment terhadap invoice',
    ]
    for item in invoicing:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.6. Gudang (Warehouse)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Warehouse" (package) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola inventory hotel termasuk amenities, supplies, linen, F&B items, '
        'dan maintenance spare parts.'
    )

    doc.add_paragraph('Modul-modul Warehouse:', style='Heading 3')

    doc.add_paragraph('A. Master Data Inventory', style='Heading 4')
    doc.add_paragraph('Cara mengakses: Warehouse → Master Data')
    master_data = [
        'Daftar semua item di inventory',
        'Kolom: Nama Item, Kategori, Stock Saat Ini, Minimum Stock, Unit, Supplier',
        'Kategori: Guest Amenities, Linen & Towels, Cleaning Supplies, F&B, Maintenance, Office Supplies',
        'Alert merah untuk item yang stock-nya di bawah minimum',
        'Tambah item baru dengan tombol "+ Add Item"',
        'Edit atau hapus item existing',
    ]
    for item in master_data:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Stock Movement', style='Heading 4')
    doc.add_paragraph('Cara mengakses: Warehouse → Stock Movements')
    stock_movement = [
        'Record semua pergerakan barang (masuk/keluar)',
        'Jenis movement:',
        '   - Stock In: Barang masuk dari supplier',
        '   - Stock Out: Barang keluar untuk digunakan',
        '   - Adjustment: Koreksi stock (stock opname)',
        '   - Transfer: Pindah antar lokasi',
        'Cara mencatat Stock In:',
        '   1. Klik "+ New Movement"',
        '   2. Pilih Movement Type: "Stock In"',
        '   3. Pilih Item',
        '   4. Masukkan Quantity',
        '   5. Pilih Supplier (opsional)',
        '   6. Tanggal',
        '   7. Notes (opsional)',
        '   8. Klik "Save"',
        'Cara mencatat Stock Out:',
        '   - Sama seperti Stock In, tapi pilih "Stock Out"',
        '   - Masukkan alasan penggunaan di Notes',
    ]
    for i, item in enumerate(stock_movement):
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('C. Purchase Orders', style='Heading 4')
    doc.add_paragraph('Cara mengakses: Warehouse → Purchase Orders')
    purchase_orders = [
        'Buat PO (Purchase Order) untuk pembelian barang',
        'Langkah membuat PO:',
        '   1. Klik "+ New Purchase Order"',
        '   2. Pilih Supplier',
        '   3. Tambahkan items yang akan dibeli',
        '   4. Masukkan quantity dan harga per unit',
        '   5. Sistem akan hitung total otomatis',
        '   6. Tanggal PO dan expected delivery',
        '   7. Klik "Create PO"',
        'Status PO:',
        '   - Draft: PO masih draft',
        '   - Sent: Sudah dikirim ke supplier',
        '   - Partial: Barang diterima sebagian',
        '   - Received: Semua barang sudah diterima',
        '   - Cancelled: PO dibatalkan',
        'Update status saat barang diterima',
        'System otomatis update inventory saat PO diterima',
    ]
    for i, item in enumerate(purchase_orders):
        if item.startswith('   -'):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        elif item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Number')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    doc.add_heading('4.7. Supplier', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Suppliers" (archive) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola database supplier/vendor yang memasok barang dan jasa ke hotel.'
    )

    doc.add_paragraph('A. Melihat Daftar Supplier', style='Heading 3')
    supplier_list = [
        'Tabel menampilkan semua supplier',
        'Kolom: Nama, Contact Person, Email, Telepon, Kategori, Status',
        'Kategori: Amenities, F&B, Linen, Maintenance, Office, Others',
        'Status: Active (hijau) atau Inactive (merah)',
        'Filter berdasarkan kategori dan status',
        'Search berdasarkan nama supplier',
    ]
    for item in supplier_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Menambah Supplier Baru', style='Heading 3')
    new_supplier = [
        'Klik tombol "+ Add Supplier"',
        'Isi form supplier:',
        '   Company Info:',
        '     - Nama perusahaan',
        '     - Kategori',
        '     - Alamat lengkap',
        '   Contact Info:',
        '     - Contact person',
        '     - Email',
        '     - Telepon',
        '     - Website (opsional)',
        '   Financial Info:',
        '     - Payment terms (NET 30, NET 45, dll)',
        '     - Bank account details',
        '     - Tax ID',
        'Klik "Save Supplier"',
    ]
    for i, item in enumerate(new_supplier, 1):
        if item.startswith('     -'):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        elif item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet')
        else:
            doc.add_paragraph(f'{item}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Melihat Riwayat Transaksi dengan Supplier', style='Heading 3')
    doc.add_paragraph(
        'Di halaman detail supplier, bagian bawah menampilkan semua purchase orders '
        'yang pernah dibuat ke supplier tersebut.'
    )

    doc.add_heading('4.8. Laporan (Reports)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Reports" (file) di Office Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Generate berbagai laporan untuk analisis bisnis dan decision making.'
    )

    doc.add_paragraph('Jenis-jenis Laporan:', style='Heading 3')

    report_types = [
        ('Occupancy Report',
         'Laporan tingkat hunian kamar per periode dengan breakdown per room type dan grafik trends.'),
        ('Revenue Report',
         'Laporan pendapatan dengan breakdown per source (room, F&B, services), comparison dengan periode sebelumnya.'),
        ('Tax Report (Laporan Pajak)',
         'Laporan pajak untuk pemerintah Indonesia dengan rincian PPN, Pajak Hotel, dan PPh Final. '
         'Update November 2025: Data individual transaksi sudah ditampilkan dengan benar.'),
        ('Guest Report',
         'Analisis tamu: nationality breakdown, repeat guests, average length of stay, guest satisfaction.'),
        ('Housekeeping Report',
         'Performa housekeeping: jumlah kamar dibersihkan, average cleaning time, staff performance.'),
        ('Inventory Report',
         'Stock position, stock movements, low stock alerts, inventory value.'),
        ('Financial Statement',
         'Income statement, balance sheet, cash flow statement.'),
        ('Employee Report',
         'Staff attendance, overtime, leave balance, payroll summary.'),
    ]

    for report_name, description in report_types:
        p = doc.add_paragraph()
        p.add_run(f'{report_name}: ').bold = True
        p.add_run(description)

    doc.add_paragraph()
    doc.add_paragraph('Detail Laporan Pajak:', style='Heading 4')
    doc.add_paragraph(
        'Laporan Pajak adalah laporan khusus untuk pelaporan pemerintah Indonesia '
        'yang mencakup seluruh transaksi hotel dengan breakdown pajak lengkap.'
    )

    tax_report_content = [
        'Informasi yang ditampilkan:',
        '   - Periode laporan (format: YYYY-MM)',
        '   - Revenue breakdown (kamar + event)',
        '   - PPN (Pajak Pertambahan Nilai): 11%',
        '   - Pajak Hotel (Pajak Daerah): 10%',
        '   - Service Charge: 10%',
        '   - PPh Final: 10% dari gross revenue',
        '   - Payment methods breakdown',
        '   - Daily revenue breakdown',
        '   - Daftar transaksi individual dengan data tamu',
        'Cara mengakses:',
        '   - API endpoint: /api/hotel/reports/tax/?period=YYYY-MM',
        '   - Contoh: /api/hotel/reports/tax/?period=2025-11',
        'Format output: JSON (dapat di-export ke Excel/PDF)',
        'Digunakan untuk: Pelaporan pajak bulanan ke pemerintah',
    ]
    for item in tax_report_content:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Update November 2025: ').bold = True
    p.add_run(
        'Sistem telah diperbaiki untuk memastikan semua transaksi individual menampilkan '
        'nilai yang benar (subtotal, pajak, service charge, grand total). Data sekarang '
        'dihitung berdasarkan pembayaran aktual yang diterima, bukan dari total reservasi.'
    )

    doc.add_paragraph()
    doc.add_paragraph('Cara Generate Laporan:', style='Heading 3')
    generate_report = [
        'Pilih jenis laporan yang diinginkan',
        'Pilih periode/date range',
        'Pilih filter tambahan (opsional): room type, department, dll',
        'Klik "Generate Report"',
        'Sistem akan menampilkan preview laporan',
        'Opsi export:',
        '   - Download PDF (untuk print atau arsip)',
        '   - Download Excel (untuk analisis lebih lanjut)',
        '   - Send via Email (kirim langsung ke email)',
    ]
    for i, step in enumerate(generate_report, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_heading('4.9. Pengaturan Sistem (Settings)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Office Settings" (gear) di Office Sidebar bagian bawah')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengkonfigurasi berbagai pengaturan sistem hotel. Hanya accessible oleh Admin dan Management.'
    )

    doc.add_paragraph('Kategori Pengaturan:', style='Heading 3')

    settings_categories = [
        ('General Settings', [
            'Nama hotel',
            'Logo hotel (upload)',
            'Alamat lengkap',
            'Contact info (telepon, email, website)',
            'Timezone',
            'Currency',
            'Date format',
        ]),
        ('Room Settings', [
            'Manage room types (add, edit, delete)',
            'Room rates per room type',
            'Room amenities',
            'Extra bed charges',
            'Child policy',
        ]),
        ('Booking Settings', [
            'Booking confirmation mode: Manual/Auto',
            'Cancellation policy',
            'No-show policy',
            'Prepayment requirements',
            'Check-in time',
            'Check-out time',
        ]),
        ('Payment Settings', [
            'Accepted payment methods',
            'Tax rate (%)',
            'Service charge (%)',
            'Payment gateway configuration (jika ada)',
        ]),
        ('Email Settings', [
            'SMTP configuration',
            'Email templates (booking confirmation, invoice, dll)',
            'Automated email triggers',
        ]),
        ('User & Permissions', [
            'Manage user roles',
            'Permission settings per role',
            'Department access configuration',
        ]),
    ]

    for category_name, settings_list in settings_categories:
        doc.add_paragraph(category_name, style='Heading 4')
        for setting in settings_list:
            doc.add_paragraph(setting, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Chapter 5: Support Operations
    doc.add_heading('5. SUPPORT - HOUSEKEEPING & MAINTENANCE', 1)

    doc.add_heading('5.1. Support Dashboard', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph(
        'Setelah login sebagai staff housekeeping/maintenance, Anda akan langsung masuk ke Support Dashboard. '
        'Atau klik logo hotel di sidebar untuk kembali ke dashboard.'
    )

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    support_dashboard = [
        'Total tugas hari ini (housekeeping + maintenance)',
        'Tugas pending yang perlu segera dikerjakan',
        'Tugas in progress yang sedang dikerjakan',
        'Tugas completed hari ini',
        'Permintaan amenities dari tamu yang perlu dipenuhi',
        'Alert kamar yang perlu perhatian khusus',
        'Grafik performa harian/mingguan',
    ]
    for item in support_dashboard:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2. Maintenance', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Maintenance" (wrench) di Support Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola semua work order maintenance, baik preventive maintenance maupun '
        'repair request dari staff atau tamu.'
    )

    doc.add_paragraph('A. Melihat Daftar Maintenance Tasks', style='Heading 3')
    maintenance_list = [
        'Tabel menampilkan semua maintenance tasks',
        'Kolom: Work Order #, Kamar/Lokasi, Kategori, Prioritas, Status, Assigned To, Created Date',
        'Filter berdasarkan status: Pending, In Progress, Completed',
        'Filter berdasarkan prioritas: Low, Medium, High, Urgent',
        'Filter berdasarkan kategori: Electrical, Plumbing, HVAC, Furniture, Appliances, Others',
        'Badge merah untuk urgent tasks',
    ]
    for item in maintenance_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Membuat Maintenance Task Baru', style='Heading 3')
    new_maintenance = [
        'Klik tombol "+ New Maintenance Task"',
        'Isi form:',
        '   - Lokasi: Pilih kamar atau area (Lobby, Pool, Gym, dll)',
        '   - Kategori: Electrical/Plumbing/HVAC/dll',
        '   - Deskripsi masalah: Jelaskan detail',
        '   - Prioritas: Low/Medium/High/Urgent',
        '   - Upload foto (jika ada)',
        'Klik "Create Task"',
        'Task tersimpan dengan status "Pending"',
        'Maintenance staff akan menerima notifikasi',
    ]
    for i, step in enumerate(new_maintenance, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Mengerjakan Maintenance Task', style='Heading 3')
    work_maintenance = [
        'Klik pada task yang akan dikerjakan',
        'Klik "Start Work" untuk mengubah status menjadi "In Progress"',
        'Sistem akan record start time otomatis',
        'Setelah selesai:',
        '   - Klik "Complete Task"',
        '   - Isi form completion:',
        '     • Deskripsi pekerjaan yang dilakukan',
        '     • Parts/materials yang digunakan',
        '     • Upload foto hasil (before/after)',
        '   - Klik "Submit"',
        'Status berubah menjadi "Completed"',
        'Sistem record completion time dan hitung duration',
        'Jika perlu follow-up, buat task baru yang terhubung',
    ]
    for i, step in enumerate(work_maintenance):
        if step.startswith('     •'):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        elif step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet')
        else:
            doc.add_paragraph(step, style='List Number')

    doc.add_heading('5.3. Housekeeping', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Housekeeping" (circular reload) di Support Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola tugas pembersihan kamar. Staff housekeeping menggunakan halaman ini '
        'untuk melihat room assignment dan update status cleaning.'
    )

    doc.add_paragraph('A. Melihat Room Assignment', style='Heading 3')
    housekeeping_view = [
        'Tabel menampilkan semua kamar yang perlu dibersihkan',
        'Prioritas:',
        '   - Kamar "Dirty" (setelah check-out) - prioritas tertinggi',
        '   - Kamar "Occupied" yang perlu daily cleaning',
        '   - Kamar "Available" yang perlu inspection',
        'Kolom: Nomor Kamar, Status Kamar, Tipe, Prioritas, Assigned To, Notes',
        'Filter berdasarkan floor/lantai',
        'Filter berdasarkan status: Dirty, Cleaning, Clean',
        'Warna badge menunjukkan prioritas:',
        '   - Merah: Urgent (check-in scheduled soon)',
        '   - Kuning: Medium',
        '   - Hijau: Low',
    ]
    for i, item in enumerate(housekeeping_view):
        if item.startswith('   -'):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        elif item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Mulai Membersihkan Kamar', style='Heading 3')
    start_cleaning = [
        'Klik pada kamar yang akan dibersihkan',
        'Klik "Start Cleaning"',
        'Status kamar berubah menjadi "Cleaning"',
        'Sistem record start time',
        'Badge oranye muncul di room card',
    ]
    for step in start_cleaning:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Checklist Pembersihan', style='Heading 3')
    doc.add_paragraph(
        'Saat membersihkan kamar, pastikan semua item di checklist dikerjakan:'
    )
    cleaning_checklist = [
        'Bedroom:',
        '   - Make bed dengan linen bersih',
        '   - Vacuum carpet/mop floor',
        '   - Dust furniture',
        '   - Check & replace amenities (toiletries)',
        '   - Empty trash bins',
        'Bathroom:',
        '   - Clean toilet, sink, shower',
        '   - Replace towels',
        '   - Restock toilet paper, tissues',
        '   - Clean mirror',
        '   - Mop floor',
        'General:',
        '   - Check AC/lights/TV working properly',
        '   - Check minibar (jika ada)',
        '   - Arrange furniture',
        '   - Final inspection',
    ]
    for item in cleaning_checklist:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('D. Menyelesaikan Cleaning', style='Heading 3')
    complete_cleaning = [
        'Setelah selesai membersihkan, klik "Mark as Clean"',
        'Checklist akan muncul, pastikan semua item dicentang',
        'Tambahkan notes jika ada (contoh: "AC remote rusak")',
        'Upload foto hasil cleaning (opsional)',
        'Klik "Submit"',
        'Status kamar berubah menjadi "Available" (siap ditempati)',
        'Sistem record completion time',
        'Room card berubah hijau',
    ]
    for step in complete_cleaning:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('E. Melaporkan Masalah', style='Heading 3')
    doc.add_paragraph(
        'Jika menemukan masalah saat membersihkan (AC rusak, kran bocor, dll):'
    )
    report_issue = [
        'Di halaman detail kamar, klik "Report Issue"',
        'Pilih kategori masalah: Maintenance required',
        'Jelaskan detail masalah',
        'Upload foto',
        'Klik "Submit Report"',
        'Maintenance task akan otomatis dibuat',
        'Maintenance team akan menerima notifikasi',
    ]
    for step in report_issue:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('5.4. Permintaan Amenities (Amenities Request)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Amenities Request" (package) di Support Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Mengelola permintaan amenities atau barang tambahan dari tamu '
        '(extra towels, pillows, toiletries, dll).'
    )

    doc.add_paragraph('A. Melihat Daftar Permintaan', style='Heading 3')
    amenities_list = [
        'Tabel menampilkan semua permintaan amenities',
        'Kolom: Request #, Kamar, Nama Tamu, Item, Quantity, Status, Requested Date',
        'Filter berdasarkan status: Pending, In Progress, Completed',
        'Badge kuning untuk pending requests (perlu segera diproses)',
        'Badge biru untuk in progress',
        'Badge hijau untuk completed',
    ]
    for item in amenities_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('B. Membuat Permintaan Baru (dari staff)', style='Heading 3')
    doc.add_paragraph(
        'Jika tamu menelepon atau meminta langsung ke staff housekeeping:'
    )
    new_amenities = [
        'Klik tombol "+ New Request"',
        'Isi form:',
        '   - Pilih Kamar/Tamu',
        '   - Pilih Item yang diminta (dari dropdown)',
        '   - Quantity',
        '   - Priority: Normal/Urgent',
        '   - Special notes (opsional)',
        'Klik "Create Request"',
        'Request tersimpan dengan status "Pending"',
    ]
    for i, step in enumerate(new_amenities, 1):
        if step.startswith('   '):
            doc.add_paragraph(step.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(f'{step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('C. Memproses Permintaan', style='Heading 3')
    process_amenities = [
        'Klik pada request yang akan diproses',
        'Klik "Start Processing"',
        'Status berubah menjadi "In Progress"',
        'Ambil barang dari gudang/storage',
        'Antar ke kamar tamu',
        'Klik "Mark as Completed"',
        'Confirm delivery dengan tamu (signature jika perlu)',
        'Klik "Submit"',
        'Status berubah menjadi "Completed"',
        'Inventory otomatis ter-update (stock berkurang)',
    ]
    for step in process_amenities:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('5.5. Laporan Support (Support Reports)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Reports" (file) di Support Sidebar')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Generate laporan performa support team untuk evaluasi dan analisis.'
    )

    doc.add_paragraph('Jenis Laporan Support:', style='Heading 3')
    support_reports = [
        'Housekeeping Performance Report:',
        '   - Total kamar dibersihkan per staff',
        '   - Average cleaning time',
        '   - Quality inspection results',
        'Maintenance Report:',
        '   - Total tasks completed',
        '   - Average response time',
        '   - Breakdown per kategori (electrical, plumbing, dll)',
        'Amenities Usage Report:',
        '   - Item paling banyak diminta',
        '   - Cost per room',
        '   - Inventory consumption trends',
    ]
    for item in support_reports:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].bold = True

    doc.add_page_break()

    # Chapter 6: Common Features
    doc.add_heading('6. FITUR UMUM', 1)

    doc.add_heading('6.1. Profil Pengguna (Profile)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Profile" (user) di sidebar bagian bawah')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Melihat dan mengedit informasi profil Anda sendiri.'
    )

    doc.add_paragraph('Informasi yang ditampilkan:', style='Heading 3')
    profile_info = [
        'Personal Information:',
        '   - Foto profil',
        '   - Nama lengkap',
        '   - Email',
        '   - Nomor telepon',
        '   - Tanggal lahir',
        'Employment Information:',
        '   - Employee ID',
        '   - Departemen',
        '   - Jabatan',
        '   - Tanggal bergabung',
        'Account Information:',
        '   - Email login',
        '   - Last login time',
        '   - Account status',
    ]
    for item in profile_info:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('Mengedit Profil:', style='Heading 3')
    edit_profile = [
        'Klik tombol "Edit Profile"',
        'Update informasi yang bisa diubah (nama, telepon, foto)',
        'Klik "Save Changes"',
        'Catatan: Email login dan data employment hanya bisa diubah oleh Admin',
    ]
    for step in edit_profile:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Mengubah Password:', style='Heading 3')
    change_password = [
        'Klik tab "Security"',
        'Klik "Change Password"',
        'Masukkan password lama',
        'Masukkan password baru',
        'Konfirmasi password baru',
        'Klik "Update Password"',
        'Password requirements: minimal 8 karakter, kombinasi huruf dan angka',
    ]
    for step in change_password:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('6.2. Kalender (Calendar)', 2)

    doc.add_paragraph('Cara mengakses:', style='Heading 3')
    doc.add_paragraph('Klik ikon "Calendar" di sidebar bagian bawah')

    doc.add_paragraph('Fungsi:', style='Heading 3')
    doc.add_paragraph(
        'Melihat dan mengelola semua event, schedule, dan deadline dalam format kalender.'
    )

    doc.add_paragraph('Fitur Kalender:', style='Heading 3')
    calendar_features = [
        'View mode: Month, Week, Day',
        'Color-coded events berdasarkan kategori:',
        '   - Biru: Booking/Check-in/Check-out',
        '   - Hijau: Meeting/Event',
        '   - Merah: Maintenance urgent',
        '   - Ungu: Training/Staff event',
        '   - Oranye: Deadline/Important date',
        'Klik event untuk melihat detail',
        'Add new event dengan klik tanggal kosong',
        'Sync dengan personal calendar (Google Calendar, Outlook)',
        'Set reminder untuk event penting',
    ]
    for item in calendar_features:
        if item.startswith('   '):
            doc.add_paragraph(item.strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3. Logout', 2)

    doc.add_paragraph('Cara logout:', style='Heading 3')
    logout_steps = [
        'Klik ikon "Logout" (merah) di pojok kanan atas navbar',
        'Konfirmasi jika diminta',
        'Anda akan diarahkan kembali ke halaman login',
        'Session Anda akan dihapus dari sistem',
    ]
    for step in logout_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Penting:', style='Heading 3')
    logout_notes = [
        'Selalu logout setelah selesai menggunakan sistem, terutama di komputer bersama',
        'Jika lupa logout, session akan otomatis expire setelah tidak aktif selama 24 jam',
        'Jangan share kredensial login Anda dengan orang lain',
    ]
    for note in logout_notes:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_page_break()

    # Chapter 7: Tips & Troubleshooting
    doc.add_heading('7. TIPS & TROUBLESHOOTING', 1)

    doc.add_heading('Tips Penggunaan Sistem', 2)

    tips = [
        ('Gunakan Search',
         'Hampir setiap halaman memiliki search bar. Gunakan untuk menemukan data dengan cepat '
         'daripada scroll manual.'),
        ('Perhatikan Badge Notifikasi',
         'Badge merah di sidebar menunjukkan item yang perlu perhatian. Check secara berkala.'),
        ('Manfaatkan Filter',
         'Gunakan filter untuk menyaring data sesuai kebutuhan (status, tanggal, kategori).'),
        ('Backup Data Penting',
         'Untuk laporan penting, selalu download/export sebagai backup.'),
        ('Keyboard Shortcuts',
         'Ctrl+S: Save form | Ctrl+F: Search | Esc: Close modal/dialog'),
        ('Mobile Access',
         'Sistem responsive, bisa diakses dari tablet atau smartphone untuk monitoring on-the-go.'),
        ('Dark Mode',
         'Gunakan dark mode (ikon bulan/matahari) untuk mengurangi eye strain saat malam hari.'),
    ]

    for tip_title, tip_desc in tips:
        p = doc.add_paragraph()
        p.add_run(f'{tip_title}: ').bold = True
        p.add_run(tip_desc)

    doc.add_heading('Troubleshooting - Masalah Umum', 2)

    doc.add_paragraph('Masalah 1: Tidak Bisa Login', style='Heading 3')
    login_troubleshoot = [
        'Periksa: Email dan password sudah benar? (case-sensitive)',
        'Periksa: Caps Lock tidak aktif?',
        'Periksa: Koneksi internet stabil?',
        'Coba: Clear browser cache dan cookies',
        'Coba: Gunakan browser lain (Chrome, Firefox, Safari)',
        'Jika masih gagal: Hubungi administrator untuk reset password',
    ]
    for item in login_troubleshoot:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 2: Halaman Tidak Muncul / Blank', style='Heading 3')
    blank_page = [
        'Coba: Refresh halaman (F5 atau Ctrl+R)',
        'Coba: Clear browser cache (Ctrl+Shift+Delete)',
        'Coba: Logout dan login kembali',
        'Periksa: Koneksi internet',
        'Periksa: Browser versi terbaru (update jika perlu)',
    ]
    for item in blank_page:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 3: Data Tidak Tersimpan', style='Heading 3')
    data_not_saved = [
        'Periksa: Semua field required sudah diisi (bertanda *)',
        'Periksa: Format data sudah benar (email, telepon, tanggal)',
        'Periksa: Session belum expire (jika terlalu lama idle)',
        'Coba: Isi form ulang dan submit lagi',
        'Jika masih gagal: Screenshot error message dan laporkan ke IT support',
    ]
    for item in data_not_saved:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 4: Tidak Bisa Akses Menu Tertentu', style='Heading 3')
    access_denied = [
        'Kemungkinan: Anda tidak memiliki hak akses ke menu tersebut',
        'Periksa: Role dan departemen Anda di halaman Profile',
        'Contoh: Staff Housekeeping tidak bisa akses Office menu',
        'Solusi: Jika merasa harus punya akses, hubungi supervisor atau admin untuk perubahan hak akses',
    ]
    for item in access_denied:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Masalah 5: Laporan Tidak Muncul / Error', style='Heading 3')
    report_error = [
        'Periksa: Date range tidak terlalu besar (maksimal 1 tahun)',
        'Periksa: Filter yang dipilih sudah benar',
        'Coba: Gunakan date range lebih kecil',
        'Coba: Refresh dan generate ulang',
        'Jika tetap error: Hubungi IT support dengan screenshot error',
    ]
    for item in report_error:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Browser yang Direkomendasikan', 2)
    browsers = [
        'Google Chrome (versi 90+) - Recommended',
        'Mozilla Firefox (versi 88+)',
        'Safari (versi 14+) untuk Mac',
        'Microsoft Edge (versi 90+)',
    ]
    doc.add_paragraph('Sistem dioptimalkan untuk browser modern. Gunakan salah satu browser berikut:')
    for browser in browsers:
        doc.add_paragraph(browser, style='List Bullet')

    doc.add_heading('Kontak Bantuan', 2)

    contact_help = [
        ('IT Support', 'support@kapulaga-hotel.com | Ext: 101'),
        ('Admin/HR', 'admin@kapulaga-hotel.com | Ext: 102'),
        ('Management', 'management@kapulaga-hotel.com | Ext: 100'),
    ]

    doc.add_paragraph('Jika mengalami masalah atau butuh bantuan:')
    for contact_title, contact_info in contact_help:
        p = doc.add_paragraph()
        p.add_run(f'{contact_title}: ').bold = True
        p.add_run(contact_info)

    doc.add_paragraph()
    doc.add_paragraph(
        'Saat menghubungi support, siapkan informasi: Nama, Departemen, Deskripsi masalah, '
        'Screenshot (jika ada), dan langkah-langkah yang sudah dicoba.'
    )

    doc.add_page_break()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer.add_run(
        '\n\n───────────────────────────────────────\n'
        'PANDUAN PENGGUNAAN SISTEM MANAJEMEN HOTEL KAPULAGA\n'
        'Versi 1.0 - November 2025\n'
        '© 2025 Kapulaga Hotel. All rights reserved.\n'
        '───────────────────────────────────────'
    )
    footer_text.font.size = Pt(9)
    footer_text.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = '/Users/pro/Dev/ladapala/Panduan_Sistem_Hotel_Kapulaga.docx'
    doc.save(output_path)
    print(f'✓ Manual berhasil dibuat: {output_path}')
    return output_path

if __name__ == '__main__':
    create_manual()
